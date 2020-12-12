from datetime import timedelta, datetime
from os import system, name 
def clear(): 
    if name == 'nt': 
        _ = system('cls') 
    else: 
        _ = system('clear') 
import sys
import time
import colorama
from colorama import Fore, Back, Style 

import docx 
import time



#============= PROGRESS BAR ================================
# SOURCE: https://stackoverflow.com/questions/3173320/text-progress-bar-in-the-console
def printProgressBar (iteration, total, prefix = '', suffix = '', decimals = 1, length = 100, fill = '█', printEnd = "\r"):
    """
    Call in a loop to create terminal progress bar
    @params:
        iteration   - Required  : current iteration (Int)
        total       - Required  : total iterations (Int)
        prefix      - Optional  : prefix string (Str)
        suffix      - Optional  : suffix string (Str)
        decimals    - Optional  : positive number of decimals in percent complete (Int)
        length      - Optional  : character length of bar (Int)
        fill        - Optional  : bar fill character (Str)
        printEnd    - Optional  : end character (e.g. "\r", "\r\n") (Str)
    """
    percent = ("{0:." + str(decimals) + "f}").format(100 * (iteration / float(total)))
    filledLength = int(length * iteration // total)
    bar = fill * filledLength + '-' * (length - filledLength)
    print('\r%s |%s| %s%% %s' % (prefix, bar, percent, suffix), end = printEnd)
    # Print New Line on Complete
    if iteration == total: 
        print()
        
#=================================================  

#=============================ДАТА ДЕЙСТВИЯ ЛИЦЕНЗИИ
now = datetime.now()

two = timedelta(days=730)


dead = (now - two)
pred = timedelta(days=62)
dead2 = (now + pred)
#=====================================================ГЛАВНЫЙ ЭКРАН И ФУНКЦИЯ СТАРТ


                                                                                             
print("Приветствую, сегодня: " + now.strftime("%d.%m.%y"))              
print("Крайний срок равен: " + str(dead))

def age_sort():

    
    a = int(input("Введите с: "))
    b = int(input("Введите до: "))
    #=======================================ЧТЕНИЕ ФАЙЛА И ЗАПУСК ЦИКЛА ДЛЯ ПРОХОДА ПО РЯДКАХ ТАБЛИЦЫ
######
    doc = docx.Document('database.docx')  #Открытие файла
   
    table = doc.tables[0]      #Инициализация таблицы
    
    count_rows = 0 #Количество рядков

    doc_output2 = docx.Document() #Создание нового файла
    doc_output2.add_paragraph('Этот файл был создан автоматически!\nВремя создания: ' + str(now)) #а почему бы и нет
    doc_output2.add_paragraph(f"Результаты сортировки по возрасту от {a} до {b} лет") 
    
    doc_output2.add_table(1,8) #ну колонок то у нас 8
    table_output3 = doc_output2.tables[0]
    table_output3.style = "Table Grid"
    table_output3.rows[0].cells[0].text = "ФИО"
    table_output3.rows[0].cells[1].text = "Дата рождения"
    table_output3.rows[0].cells[2].text = "Тренер"
    table_output3.rows[0].cells[3].text = "Дата зачисления"
    table_output3.rows[0].cells[4].text = "Группа"
    table_output3.rows[0].cells[5].text = "Разряд"
    table_output3.rows[0].cells[6].text = "Лицензия"
    table_output3.rows[0].cells[7].text = "Приказ"
    
    
    #doc_output.save("Файл KILLIST_" + str(now.strftime("%d.%m.%y")) + ".docx")
#=======================================

    # Считает количество рядков, нужно для прогресс бара
    for row in table.rows[1:]:
        count_rows += 1 
    # Инициализация прогресс бара
    printProgressBar(0, count_rows, prefix = 'Прогресс:', suffix = 'Выполнено', length = 50)    


    i = 0  # Счетчик для прогрес бара
    time_start=time.time()  # Получаем начальное время выполнения операции 

    
#===========Цикл прохода по строчкам таблицы + отрисовка прогресс бара    
    for row in table.rows[1:]:
        
        
        birth_date = row.cells[2].text
        
        
        
        #print(f'Данные из таблицы: ФИО - {name}; Дата рождения - {birth_date}; Тренер - {trainer}; Дата зачисления - {in_date}; Группа - {group}; Разряд - {category}; Лицензия - {license}; Приказ - {order}. ') 
        
        # ОБНОВЛЕНИЕ ПРОГРЕСС БАРА
        i += 1 # Счетчик для прогрес бара
        printProgressBar(i + 1, count_rows+1, prefix = 'Прогресс:', suffix = 'Выполнено', length = 50)


    
#====================Проверка и конвертирование даты лицензии
        if birth_date == '':
            continue
        if len(birth_date) >8:
            birth_date = birth_date[0:8]
        date_object_birth = datetime.strptime(birth_date, "%d.%m.%y")
        
        delta = now - date_object_birth
        days = (delta.days)
        #над данной строчкой нужно подумать
        v = days // 365
        if v > a and v < b:
            
        
            name = row.cells[1].text
            
            trainer = row.cells[3].text
            in_date = row.cells[4].text
            group = row.cells[5].text
            category = row.cells[6].text
            license = row.cells[7].text 
            order = row.cells[8].text
            table_output3.add_row()
            table_output3.rows[-1].cells[0].text = name
            table_output3.rows[-1].cells[1].text = birth_date
            table_output3.rows[-1].cells[2].text = trainer
            table_output3.rows[-1].cells[3].text = in_date
            table_output3.rows[-1].cells[4].text = group
            table_output3.rows[-1].cells[5].text = category
            table_output3.rows[-1].cells[6].text = license
            table_output3.rows[-1].cells[7].text = order
            #print(f'Пользователь: {name} подходит под условия сортировки')
            
             
    print(f' Время выполнения операции: {int(time.time()-time_start)} секунд') #Получаем конечное время
    
    
#============Вывод сообщения и документа со списком тех, кто просрочил лицензицю
    
    doc_output2.save("Файл KILLIST_age_sortirovka" + str(now.strftime("%d.%m.%y")) + ".docx")
    print("Файл KILLIST_age_sortirovka" + str(now.strftime("%d.%m.%y")) + ".docx сохранен")
    back = input('\nНажмите "1" что-бы вернуться к меню: ')
    if back == "1":
      clear() 
      


def start():
  while True:
    while True:
      try:
        a = int(input("\n[1] для продолжения\n[4] для сортировки по возрасту \n[2] помощь\n[3] авторы\n[0] выход\n"))
      except ValueError:
        clear()
        print(Fore.BLACK + Back.RED + 'ОШИБКА!\nМОЖНО ВВОДИТЬ ТОЛЬКО ЧИСЛА!')
        print(Style.RESET_ALL)
        continue
      break

#=====================РАЗДЕЛЫ МЕНЮ

    if a == 0:
      clear() 
      print("До встречи!")
      time.sleep(1)
      sys.exit()
    elif a == 3:
      clear()
      print("Версия: 1.2\nАвторы: @nyafka666,  Bohdan Zhuravel\nДата создания: 21 мая 2020\n")
      back = input('Нажмите "1" что-бы вернуться назад: ')
      if back == "1":
        clear() 
        return
    elif a == 2:
      clear() 
      print("скоро")
      back = input('Нажмите "1" что-бы вернуться назад: ')
      if back == "1":
        clear() 
        return
    elif a == 1:
      clear() 
      licence_dead()
    elif a == 4:
      clear()
      age_sort()
      pass
    

#=======================================ЧТЕНИЕ ФАЙЛА И ЗАПУСК ЦИКЛА ДЛЯ ПРОХОДА ПО РЯДКАХ ТАБЛИЦЫ
######

def licence_dead():
    doc = docx.Document('database.docx')  #Открытие файла
   
    table = doc.tables[0]      #Инициализация таблицы
    
    count_rows = 0 #Количество рядков
#=======================================СОЗДАНИЯ ФАЙЛА И ТАБЛИЦЫ    ДЛЯ ЗАПИСИ В ТАБЛИЦУ

    doc_output = docx.Document() #Создание нового файла
    doc_output.add_paragraph('Этот файл был создан программой "KiLlLicense" автоматически!\nВремя создания: ' + str(now)) #а почему бы и нет
    doc_output.add_table(1,8) #ну колонок то у нас 8
    table_output = doc_output.tables[0]
    table_output.style = "Table Grid"########################################AAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAaa
    
    # Создание шапки таблицы
    table_output.rows[0].cells[0].text = "ФИО"
    table_output.rows[0].cells[1].text = "Дата рождения"
    table_output.rows[0].cells[2].text = "Тренер"
    table_output.rows[0].cells[3].text = "Дата зачисления"
    table_output.rows[0].cells[4].text = "Группа"
    table_output.rows[0].cells[5].text = "Разряд"
    table_output.rows[0].cells[6].text = "Лицензия"
    table_output.rows[0].cells[7].text = "Приказ"
    
    
    
    #################################################Таблица с предупреждениям
    doc_output.add_paragraph("\n Далее показаны пользователи у которых лицензия закончится в течении следующих двух месяцев!") 
    doc_output.add_table(1,8) #ну колонок то у нас 8
    table_output2 = doc_output.tables[1]
    table_output2.style = "Table Grid"
    table_output2.rows[0].cells[0].text = "ФИО"
    table_output2.rows[0].cells[1].text = "Дата рождения"
    table_output2.rows[0].cells[2].text = "Тренер"
    table_output2.rows[0].cells[3].text = "Дата зачисления"
    table_output2.rows[0].cells[4].text = "Группа"
    table_output2.rows[0].cells[5].text = "Разряд"
    table_output2.rows[0].cells[6].text = "Лицензия"
    table_output2.rows[0].cells[7].text = "Приказ"
    
    
    #doc_output.save("Файл KILLIST_" + str(now.strftime("%d.%m.%y")) + ".docx")
#=======================================

    # Считает количество рядков, нужно для прогресс бара
    for row in table.rows[1:]:
        count_rows += 1 
    # Инициализация прогресс бара
    printProgressBar(0, count_rows, prefix = 'Прогресс:', suffix = 'Выполнено', length = 50)    


    i = 0  # Счетчик для прогрес бара
    time_start=time.time()  # Получаем начальное время выполнения операции 

    
#===========Цикл прохода по строчкам таблицы + отрисовка прогресс бара    
    for row in table.rows[1:]:
        
        
        
        license = row.cells[7].text 
        
        
        #print(f'Данные из таблицы: ФИО - {name}; Дата рождения - {birth_date}; Тренер - {trainer}; Дата зачисления - {in_date}; Группа - {group}; Разряд - {category}; Лицензия - {license}; Приказ - {order}. ') 
        
        # ОБНОВЛЕНИЕ ПРОГРЕСС БАРА
        i += 1 # Счетчик для прогрес бара
        printProgressBar(i + 1, count_rows+1, prefix = 'Прогресс:', suffix = 'Выполнено', length = 50)


    
#====================Проверка и конвертирование даты лицензии
        if license == '':
            continue
        if len(license) >8:
            license = license[0:8]
        date_object = datetime.strptime(license, "%d.%m.%y")
        

        if date_object < dead:
            name = row.cells[1].text
            birth_date = row.cells[2].text
            trainer = row.cells[3].text
            in_date = row.cells[4].text
            group = row.cells[5].text
            category = row.cells[6].text
            order = row.cells[8].text
            table_output.add_row()
            table_output.rows[-1].cells[0].text = name
            table_output.rows[-1].cells[1].text = birth_date
            table_output.rows[-1].cells[2].text = trainer
            table_output.rows[-1].cells[3].text = in_date
            table_output.rows[-1].cells[4].text = group
            table_output.rows[-1].cells[5].text = category
            table_output.rows[-1].cells[6].text = license
            table_output.rows[-1].cells[7].text = order
            print(f'Для пользователя: {name} действие лицензии закончилось! Дата покупки {license}')
            #timedelta(days=62)
        else:
            if date_object + timedelta(days=730) <= dead2:
                name = row.cells[1].text
                birth_date = row.cells[2].text
                trainer = row.cells[3].text
                in_date = row.cells[4].text
                group = row.cells[5].text
                category = row.cells[6].text
                order = row.cells[8].text
                table_output2.add_row()
                table_output2.rows[-1].cells[0].text = name
                table_output2.rows[-1].cells[1].text = birth_date
                table_output2.rows[-1].cells[2].text = trainer
                table_output2.rows[-1].cells[3].text = in_date
                table_output2.rows[-1].cells[4].text = group
                table_output2.rows[-1].cells[5].text = category
                table_output2.rows[-1].cells[6].text = license
                table_output2.rows[-1].cells[7].text = order
                print(f'Для пользователя: {name} действие лицензии закончится в течении 2 месяцев! Дата покупки {license}')
             




             
    print(f' Время выполнения операции: {int(time.time()-time_start)} секунд') #Получаем конечное время
    
    
#============Вывод сообщения и документа со списком тех, кто просрочил лицензицю
    
    doc_output.save("Файл KILLIST_" + str(now.strftime("%d.%m.%y")) + ".docx")
    print("Файл KILLIST_" + str(now.strftime("%d.%m.%y")) + ".docx сохранен")
    
    
    
#========================================Возвращение к главному меню
    back = input('\nНажмите "1" что-бы вернуться к меню: ')
    if back == "1":
      clear() 
      start()

while True:
  start()